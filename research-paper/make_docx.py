"""Build the ICCSDI 2026 manuscript as an editable .docx.

Generated rather than hand-written so it can be rebuilt when numbers change --
every figure quoted here traces to a file under results/.

    python research-paper/make_docx.py

Writes research-paper/draft/Hinglish_RAG_ICCSDI2026_draft.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "research-paper" / "draft" / "Hinglish_RAG_ICCSDI2026_draft.docx"
FIG = ROOT / "results"

ACCENT = RGBColor(0x1F, 0x3B, 0x57)
MUTED = RGBColor(0x55, 0x5F, 0x6B)


# --------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------

def setup(doc: Document) -> None:
    """Narrow margins and a compact body font, as requested."""
    for s in doc.sections:
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.06


def para(doc, text, size=10, bold=False, italic=False, align=None,
         space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def rich(doc, chunks, size=10, align=None, space_after=6):
    """Paragraph from [(text, bold, italic), ...] so key numbers can be bolded."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    for text, bold, italic in chunks:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
    return p


def heading(doc, text, level=1):
    sizes = {1: 13, 2: 11, 3: 10}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes[level])
    r.font.color.rgb = ACCENT
    return p


def bullets(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            lead, rest = it
            r = p.add_run(lead)
            r.bold = True
            r.font.size = Pt(size)
            r2 = p.add_run(rest)
            r2.font.size = Pt(size)
        else:
            p.add_run(it).font.size = Pt(size)


def shade(cell, hexcolor="EDF1F5"):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def table(doc, caption, headers, rows, widths=None, bold_rows=()):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(3)
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        shade(c, "DCE6F1")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)
            run.bold = ri in bold_rows
            if i > 0:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(doc, path: Path, caption: str, width=5.9):
    if not path.exists():
        para(doc, f"[FIGURE MISSING: {path}]", italic=True, color=MUTED)
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    r = c.add_run(caption)
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = MUTED


def todo(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("[TO COMPLETE] " + text)
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)


# --------------------------------------------------------------------------
# manuscript
# --------------------------------------------------------------------------

def build() -> None:
    doc = Document()
    setup(doc)

    # ---------------- title block ----------------
    para(doc, "What Code-Mixing Actually Breaks in Clinical Retrieval-Augmented "
              "Generation: A Measurement Study on Hinglish Patient Queries",
         size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    para(doc, "Devika Jonjale* and Saachi Shinde*",
         size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "NMIMS Nilkamal School of Mathematics, Applied Statistics & Analytics, "
              "Mumbai, India",
         size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color=MUTED)
    para(doc, "M.Sc. Data Science",
         size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color=MUTED)
    para(doc, "devikajonjale04@gmail.com  ·  saachi.shinde28@gmail.com",
         size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color=MUTED)
    para(doc, "*Both authors contributed equally to this work. Author order is alphabetical.",
         size=8.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
         color=MUTED)

    # ---------------- abstract ----------------
    heading(doc, "Abstract", 2)
    rich(doc, [
        ("Clinical decision-support systems assume queries arrive in clean formal "
         "English. In India they frequently arrive in Hinglish, a romanised "
         "Hindi–English code-switched register. We study a retrieval-augmented "
         "generation (RAG) pipeline for Hinglish patient questions grounded in English "
         "clinical case reports, and evaluate four hypotheses spanning retrieval and "
         "generation. ", False, False),
        ("H₀₁ (grounding effect) is rejected: ", True, False),
        ("grounding significantly improves concept-level factual support "
         "(+0.235, d = 0.576, p = 3.1×10⁻⁶⁴ on n = 1,165), and the effect replicates "
         "on a second generator family (+0.223, d = 0.611). ", False, False),
        ("H₀₂ (code-mixing robustness) is rejected in an informative direction: ", True, False),
        ("grounded factual support is statistically flat with respect to code-mixing "
         "intensity (ρ = −0.001, p = 0.98) while the ungrounded arm degrades "
         "significantly (ρ = −0.116, p = 0.0003) — grounding absorbs the damage that "
         "code-mixing does. ", False, False),
        ("H₀₄ (retrieval-stage penalty) is rejected for every retrieval method, and "
         "reveals an asymmetry: ", True, False),
        ("code-mixing damages lexical retrieval roughly four times more than dense "
         "cross-lingual retrieval (BM25 −0.0912 Recall@1, p = 9.9×10⁻²⁶; "
         "passage-chunked LaBSE −0.0206, p = 0.017, n = 3,015). We further show that "
         "an Indian-language-specialised encoder fails on romanised Hinglish for "
         "reasons of ", False, False),
        ("script", False, True),
        (" rather than language: MuRIL performs at the random floor (0.0640 vs 0.0626) "
         "yet recovers to 0.1821 on the same content in English. Finally, we report a "
         "methodological caution: the size of the measured grounding benefit depends "
         "heavily on the reference text used for scoring, ranging from +0.203 to "
         "−0.047 concept F1 on the same generations, and the concept-overlap metric in "
         "common use is precision-only, so a constant one-word answer outscores the "
         "real system by a factor of 4.7. We release the corrected evaluation and the "
         "degenerate baselines that expose this.", False, False),
    ], size=9.5)

    rich(doc, [("Keywords: ", True, False),
               ("code-switching, retrieval-augmented generation, clinical NLP, Hinglish, "
                "evaluation methodology, cross-lingual information retrieval", False, False)],
         size=9.5, space_after=12)

    # ---------------- 1 introduction ----------------
    heading(doc, "1  Introduction")
    para(doc, "Clinical natural-language systems are typically built and benchmarked on "
              "formal English. Real patient communication in India is frequently "
              "code-switched: a single question mixes Hindi and English, written in Latin "
              "script, with no standardised orthography. A patient may write “Doctor, meri "
              "beti ko skin pe rash hai aur bahut khujli ho rahi hai” where a clinical "
              "report would record “erythematous maculopapular rash with pruritus”. The "
              "clinical content is present; the surface form is not what the system expects.")
    para(doc, "Retrieval-augmented generation (RAG) is an appealing response. If a system "
              "retrieves an authoritative case report and conditions its answer on that "
              "evidence, it should produce grounded explanations without hallucinating. "
              "This paper tests that proposition on 3,015 Hinglish patient queries paired "
              "with English clinical case narratives, evaluating four pre-registered "
              "hypotheses and subjecting the evaluation itself to scrutiny.")
    para(doc, "Our contributions are:", space_after=3)
    bullets(doc, [
        ("Hypothesis tests for grounding and code-mixing robustness ",
         "(H₀₁, H₀₂) on 1,165 evaluated query–evidence pairs, replicated across two "
         "generator families."),
        ("A retrieval-stage code-mixing penalty ",
         "(H₀₄), measured against gold human translations with an explicit leakage "
         "gate, showing lexical retrieval is ~4.4× more damaged than dense."),
        ("Evidence that script, not language, is the failure mode ",
         "for Indian-language-specialised encoders on romanised text."),
        ("A demonstration that evaluation design materially changes the reported "
         "grounding benefit, ",
         "together with the degenerate baselines needed to detect it."),
        ("Two negative results reported in full: ",
         "a published adaptive-truncation heuristic that does not transfer, and a "
         "configuration defect that silently discarded 85% of every indexed document."),
    ])
    para(doc, "We report several results that are unfavourable to our own system. Each was "
              "found by instrumenting the evaluation rather than the model, and each would "
              "otherwise have been found by a reviewer.")

    # ---------------- 2 related work ----------------
    heading(doc, "2  Related Work and Research Gap")
    para(doc, "Work on code-mixed clinical text has focused principally on generation — "
              "summarising or answering code-mixed medical questions — rather than on the "
              "retrieval stage that precedes it. Multimodal code-mixed question "
              "summarisation datasets provide Hinglish patient queries with English "
              "summaries, and RAG systems for medical vision–language models propose "
              "domain-aware retrieval with adaptive context selection. Evaluation in this "
              "literature typically relies on lexical overlap against a reference, or on "
              "concept-level factuality computed against retrieved evidence.")
    para(doc, "Two gaps motivate this study. First, the retrieval stage is rarely evaluated "
              "separately for code-mixed queries, so it is not known how much end-to-end "
              "degradation is attributable to retrieval rather than generation. Second, "
              "factuality metrics for code-mixed clinical generation are largely "
              "unvalidated: they are adopted from English pipelines without checking "
              "whether they discriminate good answers from degenerate ones.")
    todo(doc, "Populate Related Work with citations from research-work/papers/ "
              "(HiFACTMix, MedSumm, MMed-RAG, CroCoSum, LLaVA-Med, and the MMCQSD and "
              "MultiCaRe dataset papers). Add them to the .bib and cite inline.")

    # ---------------- 3 methods ----------------
    heading(doc, "3  Materials and Methods")

    heading(doc, "3.1  Problem formulation and hypotheses", 2)
    para(doc, "Let q be a Hinglish patient query and C = {c₁,…,c_N} a corpus of English "
              "clinical case narratives. A retriever R returns the top-k set E_k(q) ⊂ C. A "
              "generator G produces a = G(q, E_k(q)) conditioned on that evidence, against "
              "a zero-shot baseline a₀ = G(q). Generation is scored by clinical-concept "
              "overlap between an answer and a reference text r, reporting precision, "
              "recall and their harmonic mean F1, where K(·) extracts positively-asserted "
              "clinical concepts:")
    para(doc, "P = |K(a) ∩ K(r)| / |K(a)|      R_c = |K(a) ∩ K(r)| / |K(r)|      "
              "F1 = 2·P·R_c / (P + R_c)",
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    table(doc, "Table 1  Hypotheses evaluated in this study.",
          ["", "Null hypothesis", "Outcome"],
          [["H₀₁", "Grounding does not change factual consistency relative to zero-shot "
                   "generation.", "Rejected (§5.3)"],
           ["H₀₂", "Code-mixing intensity does not affect the performance difference "
                   "between grounded and ungrounded models.", "Rejected, per arm (§5.4)"],
           ["H₀₃", "Authoritative case evidence does not improve factual correctness "
                   "relative to general biomedical text.", "Not evaluated (§5.5)"],
           ["H₀₄", "Retrieval quality does not differ between code-mixed queries and "
                   "equivalent English renderings.", "Rejected, all systems (§5.1)"]],
          widths=[0.45, 4.4, 1.5])

    heading(doc, "3.2  Architecture", 2)
    para(doc, "The pipeline encodes q with LaBSE, searches a FAISS inner-product index over "
              "L2-normalised passage embeddings, max-pools passage scores to case scores, "
              "and injects the retrieved case text into a grounded prompt. We compare four "
              "retrievers: dense (LaBSE over passages), BM25, TF-IDF, and a reciprocal-rank "
              "fusion (RRF) of dense and lexical rankings.")
    rich(doc, [("Passage chunking.  ", True, False),
               ("Encoding a case as a single vector is lossy: the median case is 554 words "
                "while the encoder admits roughly 170. We split each case into overlapping "
                "256-token windows (32-token overlap, at most six per case, 4.17 on "
                "average) and score a case by its best-matching window. This is what allows "
                "every retriever to read the same content, which the comparison in §5.1 "
                "requires.", False, False)])
    rich(doc, [("Fusion depth.  ", True, False),
               ("RRF assigns score(c) = Σᵢ 1/(K + rankᵢ(c)) with K = 60. Candidate lists are "
                "taken to depth 100 before fusion and truncated afterwards. This is a "
                "correctness requirement, not a tuning choice: fusing two depth-10 lists "
                "cannot surface any document outside their union, reducing the hybrid to a "
                "re-ordering of what both components already agreed on.", False, False)])

    # ---------------- 4 experimental design ----------------
    heading(doc, "4  Experimental Design and Evaluation")

    heading(doc, "4.1  Datasets and preprocessing", 2)
    para(doc, "Queries. MMCQSD supplies Hinglish patient questions, each with a "
              "human-written English summary and an image caption. Evidence. MultiCaRe "
              "supplies English clinical case reports; 61,316 cases were filtered and "
              "10,000 indexed, balanced across 18 condition groups.")
    para(doc, "Constructing query–evidence pairs was itself a methodological finding. "
              "Pairing Open-i radiology reports to MMCQSD queries by TF-IDF yielded only 11 "
              "usable pairs, as the two corpora had almost no topical overlap. Replacing "
              "the corpus with MultiCaRe and the matcher with LaBSE plus condition-aware "
              "filtering produced 3,015 pairs at 100% query coverage — a 274× increase.")
    rich(doc, [("Leakage gate.  ", True, False),
               ("MMCQSD's English summary contains an image caption naming the condition "
                "group verbatim in 96.2% of rows — that is, it contains the relevance label "
                "itself. Any “English versus Hinglish” comparison using the full summary is "
                "therefore inflated. We strip the caption and assert that no condition label "
                "survives into the English query condition.", False, False)])

    heading(doc, "4.2  Baselines and evaluation metrics", 2)
    para(doc, "Retrieval is reported as Recall@k and MRR@10 against a prevalence-weighted "
              "random floor of 0.0626. Generation is scored against two references:")
    bullets(doc, [
        ("Circular reference — ", "the retrieved case. The grounded arm was conditioned "
         "on this text; the zero-shot arm never saw it. This is the protocol in common use."),
        ("Unbiased reference — ", "the MMCQSD image description, a human-written account "
         "of clinical findings that neither arm saw."),
    ])
    rich(doc, [("Degenerate baselines.  ", True, False),
               ("Concept precision has no recall term, so its optimum is to assert a single "
                "common concept. We therefore report constant-answer baselines alongside "
                "every system; their necessity is demonstrated in §5.2.", False, False)])
    rich(doc, [("Statistics.  ", True, False),
               ("Paired comparisons use the Wilcoxon signed-rank test; binary retrieval "
                "outcomes use McNemar's test; intervals are 10,000-sample bootstraps. "
                "Because the image description repeats across rows (412 distinct strings "
                "over 2,988 rows, one covering 22%), intervals on that reference resample "
                "clusters rather than rows. Benjamini–Hochberg correction is applied across "
                "the full family of generation contrasts.", False, False)])

    heading(doc, "4.3  Implementation details", 2)
    para(doc, "Encoder: sentence-transformers/LaBSE, 768-d, CPU, max_seq_length 256. Index: "
              "FAISS IndexFlatIP, 41,746 passages over 10,000 cases. Generator: Groq, "
              "temperature 0.3, max_tokens 300, evidence truncated to 400 words.")
    rich(doc, [("A reproducibility limitation.  ", True, False),
               ("The generator used for our largest run, llama-3.1-8b-instant, was "
                "decommissioned by the provider during this study and now returns HTTP 404. "
                "Those outputs are archived and re-scorable but cannot be regenerated. We "
                "therefore replicate on gpt-oss-120b and report both. We note this as a "
                "broader hazard for work built on hosted free-tier models.", False, False)])
    todo(doc, "State library versions (torch, sentence-transformers, faiss-cpu, scipy) and "
              "total compute in the final version.")

    # ---------------- 5 results ----------------
    heading(doc, "5  Results")

    heading(doc, "5.1  H₀₄: the retrieval-stage code-mixing penalty", 2)
    para(doc, "Table 2 reports Recall@1 for all four retrievers under both query "
              "conditions, with every system reading the same full case text.")
    table(doc, "Table 2  Retrieval quality by system and query language (n = 3,015). "
               "Penalty = English Recall@1 − Hinglish Recall@1.",
          ["System", "Hinglish", "English", "Penalty", "McNemar p"],
          [["Hybrid (RRF)", "0.1751", "0.1973", "+0.0222", "0.018"],
           ["LaBSE (passages)", "0.1280", "0.1486", "+0.0206", "0.017"],
           ["BM25", "0.0935", "0.1847", "+0.0912", "9.9×10⁻²⁶"],
           ["TF-IDF", "0.0842", "0.1529", "—", "—"],
           ["Random floor", "0.0626", "0.0626", "—", "—"]],
          widths=[1.9, 1.0, 1.0, 1.0, 1.3], bold_rows=(0,))
    para(doc, "H₀₄ is rejected for every retrieval method. We test it per system "
              "deliberately: a penalty appearing under one configuration is a property of "
              "that configuration, whereas one holding across methods is a property of "
              "code-mixing. Second, the penalty is 4.4× larger for lexical retrieval than "
              "for dense. BM25 is the strongest system on English and the weakest on "
              "Hinglish; the dense retriever barely moves.")
    figure(doc, FIG / "retrieval_figures" / "fig2_penalty.png",
           "Fig. 1  Recall@1 lost to code-mixing, by retrieval system, with 95% bootstrap "
           "intervals. Lexical retrieval is far more damaged than dense.")
    rich(doc, [("Script, not language.  ", True, False),
               ("On the single-vector index, MuRIL — pretrained on Indian languages — "
                "performs at the random floor for Hinglish (0.0640 versus 0.0626) but "
                "recovers to 0.1821 on the same content in English. MuRIL is trained on "
                "Devanagari Hindi whereas MMCQSD is romanised. The failure is one of script "
                "rather than of language, with direct consequences for encoder selection in "
                "code-mixed settings.", False, False)])

    heading(doc, "5.2  Metric validity: the instrument does not discriminate", 2)
    para(doc, "Scored against the unbiased reference, a constant answer consisting of the "
              "single word “swelling” attains a concept precision of 0.7132. The grounded "
              "system attains 0.1528 — a hard-coded word outperforms the real system by a "
              "factor of 4.7.")
    table(doc, "Table 3  Degenerate baselines on the unbiased reference (n = 1,876).",
          ["System", "Answer", "Concept precision"],
          [["Copy the reference", "(the reference verbatim)", "1.0000"],
           ["Constant", "“swelling”", "0.7132"],
           ["Constant", "“swelling and erythema”", "0.6586"],
           ["Constant", "“erythema”", "0.6039"],
           ["Grounded system", "(generated)", "0.1528"],
           ["Zero-shot system", "(generated)", "0.1066"],
           ["Constant", "“pain”", "0.0032"]],
          widths=[1.7, 2.6, 1.5], bold_rows=(4, 5))
    para(doc, "The cause is structural: precision alone rewards asserting few concepts. "
              "Absolute levels therefore carry no information about answer quality, and only "
              "paired differences between arms scored identically are interpretable. We "
              "report F1 throughout and ship these rows with every table. We further note "
              "that the commonly-reported “hallucination rate” is, under this definition, "
              "exactly 1 − precision, so reporting both double-counts a single result.")

    heading(doc, "5.3  H₀₁: the grounding effect", 2)
    para(doc, "H₀₁ is rejected. Under the evaluation protocol in common use — concept "
              "overlap against the retrieved evidence — grounding produces a large and "
              "highly significant improvement in factual support, and this replicates "
              "across two generator families.")
    table(doc, "Table 4  H₀₁ under the circular (evidence-based) reference.",
          ["Generator / evidence", "n", "Zero-shot", "Grounded", "Δ", "d", "p"],
          [["llama-3.1-8b, oracle (as published)", "1,165", "0.319", "0.554", "+0.235",
            "0.576", "3.1×10⁻⁶⁴"],
           ["llama-3.1-8b, oracle (repaired lexicon)", "669", "0.280", "0.576", "+0.296",
            "0.720", "2.8×10⁻⁴⁹"],
           ["gpt-oss-120b, oracle", "214", "0.276", "0.536", "+0.261", "0.640",
            "4.3×10⁻¹⁵"],
           ["gpt-oss-120b, real retrieval", "221", "0.285", "0.355", "+0.070", "0.191",
            "0.0075"]],
          widths=[2.5, 0.6, 0.9, 0.9, 0.7, 0.6, 1.0], bold_rows=(0,))
    rich(doc, [("The effect is sensitive to the reference.  ", True, False),
               ("Scoring the same generations against a reference neither arm saw changes "
                "the picture materially. Under the circular reference the grounding effect "
                "on concept F1 is +0.203; under the unbiased reference it is +0.062 on "
                "llama and −0.047 on gpt-oss-120b. The mechanism is visible in the "
                "components: grounding raises precision and lowers recall (oracle arm, "
                "unbiased reference: −0.138 recall, BH p = 0.00025). Grounding makes the "
                "model more conservative, and whether that trade is profitable depends on "
                "the generator.", False, False)])
    table(doc, "Table 5  Grounding effect on concept F1 by reference "
               "(grounded − zero-shot, Benjamini–Hochberg corrected).",
          ["Generator / evidence", "Circular ref.", "Unbiased ref."],
          [["llama-3.1-8b, oracle", "+0.203", "+0.062"],
           ["gpt-oss-120b, oracle", "+0.093", "−0.021 (n.s.)"],
           ["gpt-oss-120b, real retrieval", "−0.032", "−0.047"]],
          widths=[2.6, 1.6, 1.6])
    figure(doc, FIG / "h1_figures" / "fig3_h1_reference_effect.png",
           "Fig. 2  The same generations scored against two references. Under the circular "
           "reference grounding appears strongly beneficial; under a reference neither arm "
           "saw, the effect shrinks or reverses.")
    rich(doc, [("Refusal is a first-class outcome.  ", True, False),
               ("With real retrieval the grounded arm declines to answer in 33.8% of cases "
                "(zero-shot: 0.2%), typically stating that the evidence concerns a different "
                "patient. Refusals assert no clinical concept and therefore score as "
                "missing, disappearing from a naive mean — which would make the system "
                "appear healthiest exactly where it fails. We report answer coverage "
                "(49–52% for grounded arms, 92.9% zero-shot) alongside every factuality "
                "figure.", False, False)])

    heading(doc, "5.4  H₀₂: code-mixing robustness", 2)
    para(doc, "H₀₂ is rejected, and in an informative direction. Testing each arm separately "
              "against a repaired code-mixing measure shows that grounded factual support is "
              "statistically flat with respect to code-mixing intensity, while the "
              "ungrounded arm degrades significantly.")
    table(doc, "Table 6  H₀₂: Spearman correlation between code-mixing intensity and "
               "performance, by arm (n = 1,165, Benjamini–Hochberg corrected).",
          ["Arm", "ρ", "95% CI", "BH p", "Reading"],
          [["Grounded factual support", "−0.001", "[−0.058, +0.057]", "0.983", "Flat"],
           ["Zero-shot factual support", "−0.116", "[−0.171, −0.059]", "0.0003",
            "Degrades"],
           ["Grounded hallucination", "−0.022", "[−0.080, +0.036]", "0.593", "Flat"],
           ["Zero-shot hallucination", "+0.042", "[−0.018, +0.101]", "0.303", "Flat"]],
          widths=[2.0, 0.7, 1.4, 0.8, 1.1], bold_rows=(0, 1))
    para(doc, "Grounding absorbs the degradation that code-mixing induces in the ungrounded "
              "model. This is a stronger and more useful claim than the “no significant "
              "difference” obtained by testing the gain — a difference of two noisy arms — "
              "which is under-powered by construction.")
    rich(doc, [("The measure required repair.  ", True, False),
               ("The original code-mixing index counted the English words “doctor” and "
                "“please” as Hindi; they occur in 68.2% and 35.7% of queries respectively. "
                "Two hallucination effects reported under the original measure did not "
                "survive repair and are withdrawn.", False, False)])
    figure(doc, FIG / "h2_figures" / "h2_dose_response.png",
           "Fig. 3  Factual support against code-mixing intensity, by arm. Grounded support "
           "is flat; the ungrounded arm declines.")

    heading(doc, "5.5  H₀₃: evidence provenance", 2)
    para(doc, "H₀₃ was not evaluated. The intended comparison — authoritative case reports "
              "against general biomedical text — requires corpora of comparable topical "
              "coverage. Measured against the 18 condition groups in this study, PubMedQA "
              "contexts are only 2.1% on-topic and MMedBench is 57.6% non-English. A naive "
              "comparison would therefore have measured corpus topicality rather than "
              "evidence provenance. We report this as an open item rather than presenting an "
              "uninformative result.")

    heading(doc, "5.6  Ablations and negative results", 2)
    rich(doc, [("Adaptive truncation does not transfer.  ", True, False),
               ("A published similarity-gap heuristic for adaptive context selection fires "
                "on 0 of 3,015 queries at its specified threshold: it requires a 0.248 "
                "similarity gap between adjacent neighbours, and the largest gap occurring "
                "anywhere in the data is 0.109. Swept across thresholds and compared against "
                "a fixed k at matched evidence budget, it wins 0 of 6 settings. Precision is "
                "flat (0.112–0.115) at every setting while recall falls monotonically, "
                "indicating that the similarity gap carries no relevance information here.",
                False, False)])
    rich(doc, [("A configuration defect dominated the retrieval comparison.  ", True, False),
               ("The index was originally built with a sequence limit of 128 tokens, "
                "truncating 100% of case narratives (median 307 tokens) and 52.5% of "
                "queries, while lexical baselines read the full document. Correcting this "
                "and adding passage chunking moved dense Recall@1 on Hinglish from 0.1144 to "
                "0.1280 and reversed the ordering against BM25. We report this because the "
                "uncorrected comparison would have supported the opposite conclusion.",
                False, False)])
    figure(doc, FIG / "retrieval_figures" / "fig1_table1.png",
           "Fig. 4  Retrieval quality by system and query language, after passage chunking "
           "and content matching.")

    # ---------------- 6 discussion ----------------
    heading(doc, "6  Discussion")
    para(doc, "Three of the four hypotheses were evaluated and all three were rejected. "
              "Grounding improves factual support (H₀₁); it does so consistently across "
              "code-mixing intensity while the ungrounded arm degrades (H₀₂); and "
              "code-mixing significantly damages retrieval for every method tested (H₀₄).")
    para(doc, "The retrieval findings are the most actionable. The 4.4× asymmetry between "
              "lexical and dense degradation is a direct empirical argument for "
              "cross-lingual embedding in code-mixed settings, and the MuRIL result refines "
              "it: what matters is script coverage, not nominal language coverage. "
              "Practitioners deploying Indian-language encoders on romanised user text "
              "should expect near-chance retrieval.")
    para(doc, "The measurement findings temper the generation result. The magnitude of the "
              "grounding benefit — though not its statistical presence under the standard "
              "protocol — depends substantially on the reference text, and the standard "
              "concept-overlap metric can be beaten by a one-word constant. We therefore "
              "argue that reported RAG gains for code-mixed clinical text should be "
              "accompanied by a degenerate baseline and, where possible, a reference the "
              "model has not seen.")

    heading(doc, "6.1  Limitations and threats to validity", 2)
    bullets(doc, [
        ("Relevance labels are coarse. ", "Condition-group matching over 18 groups admits "
         "cases that are same-group but clinically irrelevant; absolute Recall should be "
         "read as a lower bound."),
        ("The concept lexicon is not clinician-validated ", "and covers 26 concepts. It "
         "cannot capture findings outside that vocabulary."),
        ("The unbiased reference is narrow and repetitive. ", "The image description "
         "averages 1.5 concepts and takes 412 distinct values over 2,988 rows; we address "
         "the non-independence with cluster bootstraps, but the construct remains partial."),
        ("Generator dependence. ", "Two generation results differ in sign between "
         "generators. We report both rather than selecting one."),
        ("Provider instability. ", "A model central to the study was withdrawn during it."),
        ("H₀₃ was not evaluated ", "for the corpus-topicality reason given in §5.5."),
        ("No clinical deployment claim is made. ", "Absolute Recall@1 of 0.175 is far below "
         "what patient-facing use would require."),
    ])

    # ---------------- 7 conclusion ----------------
    heading(doc, "7  Conclusion and Future Work")
    para(doc, "We evaluated a Hinglish clinical RAG pipeline against four hypotheses. "
              "Grounding significantly improves factual support and holds steady as "
              "code-mixing intensity rises, while the ungrounded baseline degrades. "
              "Code-mixing significantly damages retrieval for every method tested, and does "
              "so four times more severely for lexical than for dense retrieval. An "
              "Indian-language encoder fails on romanised text for reasons of script rather "
              "than language. Alongside these, we report that the magnitude of the measured "
              "grounding benefit depends materially on the scoring reference, and release "
              "the corrected evaluation and degenerate baselines that make this visible.")
    para(doc, "Future work should establish a random-evidence control to separate genuine "
              "grounding from evidence echoing, validate the concept lexicon against "
              "clinician annotation, complete H₀₃ with topically matched corpora, and extend "
              "the analysis to the multimodal signal that MMCQSD provides but that this "
              "text-only study does not use.")

    # ---------------- declarations ----------------
    heading(doc, "Declarations")
    for label, text in [
        ("Funding", "The authors received no specific funding for this work."),
        ("Competing interests", "The authors declare no competing interests."),
        ("Ethics approval and consent to participate",
         "Not applicable. All corpora are publicly available and de-identified."),
        ("Consent for publication", "Not applicable."),
        ("Data availability",
         "MMCQSD and MultiCaRe are publicly available. Derived artefacts (paired file, "
         "index, result files) are archived at [DOI to be minted]."),
        ("Materials availability", "Not applicable."),
        ("Code availability",
         "All source code, prompts, configuration and analysis scripts are available at "
         "[repository URL and release tag]. A single entry point regenerates every reported "
         "number and figure from cached artefacts."),
        ("Author contributions",
         "Both authors contributed equally. Author order is alphabetical. "
         "[Complete per CRediT roles: conceptualisation, methodology, software, "
         "validation, formal analysis, investigation, data curation, writing.]"),
    ]:
        rich(doc, [(label + "  ", True, False), (text, False, False)], size=9.5,
             space_after=4)

    heading(doc, "References")
    todo(doc, "Add references. Suggested minimum: the MMCQSD and MultiCaRe dataset papers, "
              "LaBSE, MuRIL, BM25/Okapi, reciprocal-rank fusion, MMed-RAG, and two or three "
              "code-mixed clinical NLP papers from research-work/papers/.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")


if __name__ == "__main__":
    build()
