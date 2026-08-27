"""Build the ICCSDI 2026 manuscript by FILLING THE OFFICIAL TEMPLATE.

An earlier version rebuilt the document from scratch with python-docx, which
produced a paper that did not look like the template. This one opens
`ICCSDI2026_Word_Template (1).docx` and replaces its placeholder text in place,
so every style (`ICCSDI Heading 1/2`, `ICCSDI Bullet`, `ICCSDI Caption`,
`ICCSDI Back Heading`), the Times New Roman 10pt body, the A4 page and the
title-block formatting are the template's own, not a reconstruction.

The only deliberate deviation is a small margin reduction, as requested.

Content is restricted to what the research actually established. Template
scaffolding the study does not use (the worked equation, the pseudocode block,
the optional appendix) is removed rather than filled with filler.

    python research-paper/make_paper.py

Writes research-paper/draft/Hinglish_RAG_ICCSDI2026.docx
"""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "research-paper" / "ICCSDI2026_Word_Template (1).docx"
OUT = ROOT / "research-paper" / "draft" / "Hinglish_RAG_ICCSDI2026.docx"
FIGDIR = ROOT / "results"

TITLE = ("Measuring Grounded Retrieval-Augmented Generation for Hinglish "
         "Clinical Queries")

AUTHORS = "Devika Jonjale¹* and Saachi Shinde¹*"
AFFIL = ("¹M.Sc. Data Science, NMIMS Nilkamal School of Mathematics, Applied "
         "Statistics & Analytics, Mumbai, India.")
EMAILS = ("*Corresponding author(s). E-mail(s): devikajonjale04@gmail.com; "
          "saachi.shinde28@gmail.com. Both authors contributed equally; "
          "author order is alphabetical.")

ABSTRACT = (
    "Clinical decision-support systems assume queries arrive in formal English. In "
    "India they frequently arrive in Hinglish, a romanised Hindi–English "
    "code-switched register. We evaluate a retrieval-augmented generation (RAG) "
    "pipeline for Hinglish patient questions grounded in English clinical case "
    "reports, testing four hypotheses across the retrieval and generation stages on "
    "3,015 query–evidence pairs. H₀₁ is rejected: grounding significantly improves "
    "concept-level factual support (+0.235, d = 0.576, p = 3.1×10⁻⁶⁴), and the effect "
    "replicates on a second generator family (+0.224, d = 0.616). H₀₂ is rejected in "
    "an informative direction: grounded factual support is statistically flat with "
    "respect to code-mixing intensity (ρ = −0.001, p = 0.98) while the ungrounded arm "
    "degrades significantly (ρ = −0.116, p = 0.0003), so grounding absorbs the damage "
    "code-mixing does. H₀₄ is rejected for every retrieval method tested, and reveals "
    "an asymmetry: code-mixing costs lexical retrieval 0.0912 Recall@1 "
    "(p = 9.9×10⁻²⁶) against 0.0206 for a passage-chunked dense retriever "
    "(p = 0.017). We further find that an Indian-language-specialised encoder fails "
    "on romanised Hinglish for reasons of script rather than language, performing at "
    "the random floor (0.0640 versus 0.0626) yet recovering to 0.1821 on the same "
    "content in English. Finally we report a measurement caution: the size of the "
    "grounding effect depends substantially on the reference text used for scoring, "
    "and the concept-overlap metric in common use is precision-only, so a constant "
    "one-word answer outscores the system by a factor of 4.7."
)

KEYWORDS = ("Keywords: code-switching, retrieval-augmented generation, clinical NLP, "
            "Hinglish, cross-lingual information retrieval, evaluation methodology")

# --------------------------------------------------------------------------
# Section bodies. Keyed by the template heading they follow.
# --------------------------------------------------------------------------

INTRO = [
    "Clinical natural-language systems are typically built and benchmarked on formal "
    "English. Real patient communication in India is frequently code-switched: a "
    "single question mixes Hindi and English in Latin script with no standardised "
    "orthography. A patient may write “Doctor, meri beti ko skin pe rash hai aur bahut "
    "khujli ho rahi hai” where a clinical report would record “erythematous "
    "maculopapular rash with pruritus”. The clinical content is present; the surface "
    "form is not what the system expects.",

    "Retrieval-augmented generation is an appealing response: if a system retrieves an "
    "authoritative case report and conditions its answer on that evidence, it should "
    "produce grounded explanations without hallucinating. This paper tests that "
    "proposition on 3,015 Hinglish patient queries paired with English clinical case "
    "narratives, and subjects the evaluation itself to the same scrutiny as the system.",
]

INTRO_BULLETS = [
    "Hypothesis tests for the grounding effect (H₀₁) and code-mixing robustness "
    "(H₀₂) on 1,165 evaluated pairs, replicated across two generator families.",
    "A retrieval-stage code-mixing penalty (H₀₄) measured against gold human "
    "translations with an explicit leakage gate, showing lexical retrieval is "
    "approximately 4.4 times more damaged than dense cross-lingual retrieval.",
    "Evidence that the evaluation protocol materially changes the reported grounding "
    "benefit, together with the degenerate baselines required to detect it.",
]

RELATED = [
    "Work on code-mixed clinical text has focused principally on generation — "
    "summarising or answering code-mixed medical questions — rather than on the "
    "retrieval stage that precedes it. Multimodal code-mixed question summarisation "
    "datasets supply Hinglish patient queries with English summaries, and RAG systems "
    "for medical vision–language models propose domain-aware retrieval with adaptive "
    "context selection. Evaluation in this literature typically relies on lexical "
    "overlap against a reference, or on concept-level factuality computed against the "
    "retrieved evidence.",

    "Two gaps motivate this study. First, the retrieval stage is rarely evaluated "
    "separately for code-mixed queries, so it is not known how much of the end-to-end "
    "degradation is attributable to retrieval rather than to generation. Second, "
    "factuality metrics for code-mixed clinical generation are largely unvalidated: "
    "they are adopted from English pipelines and applied without checking whether they "
    "discriminate good answers from degenerate ones. This paper addresses both.",

    "[TO COMPLETE] Add citations for the MMCQSD and MultiCaRe datasets, LaBSE, MuRIL, "
    "BM25, reciprocal-rank fusion, and the code-mixed clinical NLP work reviewed above.",
]

METHODS = [
    "The system encodes a Hinglish query with LaBSE, searches a FAISS inner-product "
    "index over L2-normalised passage embeddings, max-pools passage scores to case "
    "scores, and injects the retrieved case text into a grounded prompt. A zero-shot "
    "arm answers the same question with no evidence."
]

PROBLEM = [
    "Let q be a Hinglish patient query and C a corpus of English clinical case "
    "narratives. A retriever returns the top-k set Eₖ(q) ⊂ C, and a generator produces "
    "a = G(q, Eₖ(q)) conditioned on that evidence, against a zero-shot baseline "
    "a₀ = G(q). Generation is scored by clinical-concept overlap between an answer and "
    "a reference text r, where K(·) extracts positively-asserted clinical concepts: "
    "precision P = |K(a) ∩ K(r)| / |K(a)|, recall R = |K(a) ∩ K(r)| / |K(r)|, and their "
    "harmonic mean F₁. The choice of r is the central methodological variable of this "
    "study and is treated explicitly in Section 4.2.",

    "Four hypotheses are evaluated. H₀₁: grounding does not change factual consistency "
    "relative to zero-shot generation. H₀₂: code-mixing intensity does not affect the "
    "difference between grounded and ungrounded models. H₀₃: authoritative case "
    "evidence does not improve factual correctness relative to general biomedical "
    "text. H₀₄: retrieval quality does not differ between code-mixed queries and "
    "equivalent English renderings.",
]

ARCH = [
    "Encoding a case as a single vector is lossy: the median case is 554 words while "
    "the encoder admits roughly 170. Each case is therefore split into overlapping "
    "256-token windows (32-token overlap, at most six per case, 4.17 on average) and "
    "scored by its best-matching window. This allows every retriever — dense, BM25, "
    "TF-IDF and their reciprocal-rank fusion — to read the same content, which the "
    "comparison in Section 5.1 requires.",

    "Fusion takes candidate lists to depth 100 before combining and truncates "
    "afterwards. This is a correctness requirement rather than a tuning choice: fusing "
    "two depth-10 lists cannot surface any document outside their union, which reduces "
    "the hybrid to a re-ordering of what its components already agreed on.",
]

EXPER = [
    "All experiments use the same 3,015 query–evidence pairs, the same prompts and the "
    "same scoring, varying only the factor under test."
]

DATASETS = [
    "Queries come from MMCQSD, which supplies Hinglish patient questions each with a "
    "human-written English summary and an image caption. Evidence comes from "
    "MultiCaRe clinical case reports; 61,316 cases were filtered and 10,000 indexed, "
    "balanced across 18 condition groups.",

    "Constructing the pairs was itself a methodological finding. Matching Open-i "
    "radiology reports to MMCQSD queries by TF-IDF yielded only 11 usable pairs, as "
    "the two corpora had almost no topical overlap. Replacing the corpus with "
    "MultiCaRe and the matcher with LaBSE plus condition-aware filtering produced "
    "3,015 pairs at 100% query coverage, a 274-fold increase.",

    "MMCQSD's English summary contains an image caption naming the condition group "
    "verbatim in 96.2% of rows — that is, it contains the relevance label itself. Any "
    "English-versus-Hinglish comparison using the full summary is therefore inflated. "
    "The caption is stripped and a leakage gate asserts that no condition label "
    "survives into the English query condition.",
]

METRICS = [
    "Retrieval is reported as Recall@k and MRR@10 against a prevalence-weighted random "
    "floor of 0.0626. Generation is scored against two references: the retrieved "
    "evidence, which the grounded arm was conditioned on and the zero-shot arm never "
    "saw; and the MMCQSD image description, which neither arm saw. The first is the "
    "protocol in common use and is circular by construction; the second is unbiased.",

    "Concept precision has no recall term, so its optimum is to assert a single common "
    "concept. Constant-answer baselines are therefore reported alongside every system; "
    "their necessity is demonstrated in Section 5.2. Paired comparisons use the "
    "Wilcoxon signed-rank test, binary retrieval outcomes use McNemar's test, and "
    "intervals are 10,000-sample bootstraps. Because the image description repeats "
    "across rows (412 distinct values over 2,988 rows), intervals on that reference "
    "resample clusters rather than rows. Benjamini–Hochberg correction is applied "
    "across the full family of generation contrasts.",
]

IMPL = [
    "Encoder: sentence-transformers/LaBSE, 768-dimensional, CPU, maximum sequence "
    "length 256. Index: FAISS IndexFlatIP over 41,746 passages from 10,000 cases. "
    "Generators: Groq llama-3.1-8b-instant and openai/gpt-oss-120b, temperature 0.3, "
    "300 maximum output tokens, evidence truncated to 400 words. "
    "[TO COMPLETE] State library versions and total compute.",

    "One reproducibility limitation must be recorded. The generator used for the "
    "largest generation run, llama-3.1-8b-instant, was decommissioned by the provider "
    "during this study and now returns HTTP 404. Those outputs are archived and can be "
    "re-scored but cannot be regenerated. Results are therefore replicated on "
    "gpt-oss-120b and both are reported. This is a broader hazard for work built on "
    "hosted models.",
]

RES_INTRO = ["Results are presented in hypothesis order."]

RES_H4 = [
    "Table 1 reports Recall@1 for four retrieval systems under both query conditions, "
    "with every system reading the same full case text. H₀₄ is rejected for every "
    "method. Testing it per system is deliberate: a penalty appearing under one "
    "configuration is a property of that configuration, whereas one holding across "
    "methods is a property of code-mixing.",

    "The penalty is 4.4 times larger for lexical retrieval than for dense. BM25 is the "
    "strongest system on English and the weakest on Hinglish, while the dense "
    "retriever barely moves. On the single-vector index MuRIL, pretrained on Indian "
    "languages, performs at the random floor for Hinglish (0.0640 against a floor of "
    "0.0626) but recovers to 0.1821 on the same content in English. MuRIL is trained "
    "on Devanagari Hindi whereas MMCQSD is romanised, so the failure is one of script "
    "rather than of language.",
]

RES_METRIC = [
    "Scored against the unbiased reference, a constant answer consisting of the single "
    "word “swelling” attains a concept precision of 0.7132 while the grounded system "
    "attains 0.1528 — a hard-coded word outperforms the real system by a factor of "
    "4.7. The cause is structural: precision alone rewards asserting few concepts. "
    "Absolute levels therefore carry no information about answer quality and only "
    "paired differences between arms scored identically are interpretable. We report "
    "F₁ throughout. We further note that the commonly reported hallucination rate is, "
    "under this definition, exactly one minus precision, so reporting both "
    "double-counts a single result.",
]

RES_H1 = [
    "H₀₁ is rejected. Under the evaluation protocol in common use, grounding produces a "
    "large and highly significant improvement in factual support that replicates "
    "across two generator families: +0.235 (d = 0.576, p = 3.1×10⁻⁶⁴, n = 1,165) on "
    "llama-3.1-8b, and +0.224 (d = 0.616, p = 6.5×10⁻¹⁷, n = 237) on gpt-oss-120b under "
    "real retrieval.",

    "The magnitude is sensitive to the reference. Scoring the same generations against "
    "text neither arm saw gives +0.062 concept F₁ on llama and −0.047 on gpt-oss-120b, "
    "against +0.203 and +0.093 respectively under the circular reference. The "
    "mechanism is visible in the components: grounding raises precision and lowers "
    "recall (−0.138 recall, BH p = 0.00025), making the model more conservative, and "
    "whether that trade is profitable depends on the generator.",

    "Refusal is a first-class outcome. With real retrieval the grounded arm declines "
    "to answer in 33.8% of cases against 0.2% for zero-shot, typically stating that "
    "the evidence concerns a different patient. Refusals assert no clinical concept "
    "and therefore score as missing, disappearing from a naive mean — which would make "
    "the system appear healthiest exactly where it fails. Answer coverage is reported "
    "alongside every factuality figure.",
]

RES_H2 = [
    "H₀₂ is rejected, and in an informative direction. Testing each arm separately "
    "against a repaired code-mixing measure shows grounded factual support is "
    "statistically flat with respect to code-mixing intensity (ρ = −0.001, 95% CI "
    "[−0.058, +0.057], BH p = 0.98) while the ungrounded arm degrades significantly "
    "(ρ = −0.116, 95% CI [−0.171, −0.059], BH p = 0.0003). Grounding absorbs the "
    "degradation that code-mixing induces in the ungrounded model.",

    "This is a stronger claim than the non-significant result obtained by testing the "
    "gain, which is a difference of two noisy arms and under-powered by construction. "
    "The measure required repair: the original code-mixing index counted the English "
    "words “doctor” and “please” as Hindi, and they occur in 68.2% and 35.7% of queries "
    "respectively. Two hallucination effects reported under the original measure did "
    "not survive repair and are withdrawn.",
]

RES_H3 = [
    "H₀₃ was evaluated on four topically matched, equal-sized corpora of 1,872 "
    "documents each: MultiCaRe case reports, PubMedQA abstracts, MMedBench English "
    "exam text, and a sentence-shuffled MultiCaRe control. Matching is necessary "
    "because the corpora differ enormously in coverage of the 18 conditions studied "
    "(67.9%, 2.1% and 16.4% respectively), so an unmatched comparison would measure "
    "corpus topicality rather than evidence provenance.",

    "Evidence type significantly changes how often the model refuses to answer "
    "(Cochran's Q = 9.09, df = 3, p = 0.028, n = 160). Case reports — the "
    "authoritative source — provoke the most refusals at 88.1%, against 76.2% for exam "
    "text. Answer quality remains undetermined: pairwise differences are small and "
    "none significant, and the four-way omnibus rests on only 13 rows where all four "
    "conditions produced a scoreable answer, because refusal rates of 76–88% make that "
    "joint event rare. We report this as absence of evidence rather than evidence of "
    "absence.",
]

DISCUSSION = [
    "Three of the four hypotheses were evaluated and rejected. Grounding improves "
    "factual support; it does so consistently across code-mixing intensity while the "
    "ungrounded arm degrades; and code-mixing significantly damages retrieval for "
    "every method tested. H₀₃ is answered for refusal behaviour and open for answer "
    "quality.",

    "The retrieval findings are the most actionable. The 4.4-fold asymmetry between "
    "lexical and dense degradation is a direct empirical argument for cross-lingual "
    "embedding in code-mixed settings, and the MuRIL result refines it: what matters "
    "is script coverage rather than nominal language coverage. Practitioners deploying "
    "Indian-language encoders on romanised user text should expect near-chance "
    "retrieval.",

    "The measurement findings temper the generation result. The magnitude of the "
    "grounding benefit, though not its statistical presence under the standard "
    "protocol, depends substantially on the reference text, and the standard "
    "concept-overlap metric can be beaten by a one-word constant. Reported RAG gains "
    "for code-mixed clinical text should therefore be accompanied by a degenerate "
    "baseline and, where possible, a reference the model has not seen.",
]

LIMITS = [
    "Relevance labels are coarse: condition-group matching over 18 groups admits cases "
    "that are same-group but clinically irrelevant, so absolute Recall should be read "
    "as a lower bound. The concept lexicon covers 26 concepts and is not "
    "clinician-validated. The unbiased reference is narrow and repetitive, averaging "
    "1.5 concepts and taking 412 distinct values over 2,988 rows; non-independence is "
    "addressed with cluster bootstraps but the construct remains partial. Two "
    "generation results differ in sign between generators and both are reported rather "
    "than one selected. A model central to the study was withdrawn by its provider "
    "during the work. H₀₃'s answer-quality arm is under-powered for the reason given "
    "in Section 5.5. Finally, an absolute Recall@1 of 0.175 is far below what "
    "patient-facing deployment would require, and no clinical deployment claim is made.",
]

CONCLUSION = [
    "We evaluated a Hinglish clinical RAG pipeline against four hypotheses. Grounding "
    "significantly improves factual support and holds steady as code-mixing intensity "
    "rises while the ungrounded baseline degrades. Code-mixing significantly damages "
    "retrieval for every method tested, and four times more severely for lexical than "
    "for dense retrieval. An Indian-language encoder fails on romanised text for "
    "reasons of script rather than language. Evidence type changes how often a "
    "grounded model declines to answer, though its effect on answer quality remains "
    "undetermined.",

    "Alongside these we report that the measured size of the grounding benefit depends "
    "materially on the scoring reference, and release the corrected evaluation and the "
    "degenerate baselines that make this visible. Future work should establish a "
    "random-evidence control to separate genuine grounding from evidence echoing, "
    "validate the concept lexicon against clinician annotation, and complete the "
    "provenance comparison with a prompt that refuses less often.",
]

DECLARATIONS = [
    ("Funding", "The authors received no specific funding for this work."),
    ("Competing interests", "The authors declare no competing interests."),
    ("Ethics approval and consent to participate",
     "Not applicable. All corpora are publicly available and de-identified."),
    ("Consent for publication", "Not applicable."),
    ("Data availability",
     "MMCQSD and MultiCaRe are publicly available. Derived artefacts and all result "
     "files are archived at [TO COMPLETE: repository DOI]."),
    ("Materials availability", "Not applicable."),
    ("Code availability",
     "All source code, prompts, configuration and analysis scripts are available at "
     "[TO COMPLETE: repository URL and release tag]. A single entry point regenerates "
     "every reported number and figure from cached artefacts."),
    ("Author contributions",
     "Both authors contributed equally to conceptualisation, methodology, software, "
     "formal analysis, investigation and writing. Author order is alphabetical."),
]

TABLE1 = (
    "Table 1  Retrieval quality by system and query language (n = 3,015). Penalty is "
    "English Recall@1 minus Hinglish Recall@1.",
    ["System", "Hinglish", "English", "Penalty", "McNemar p"],
    [["Hybrid (RRF)", "0.1751", "0.1973", "+0.0222", "0.018"],
     ["LaBSE (passages)", "0.1280", "0.1486", "+0.0206", "0.017"],
     ["BM25", "0.0935", "0.1847", "+0.0912", "9.9×10⁻²⁶"],
     ["TF-IDF", "0.0842", "0.1529", "—", "—"],
     ["Random floor", "0.0626", "0.0626", "—", "—"]],
)

TABLE2 = (
    "Table 2  Grounding effect on concept F₁ (grounded minus zero-shot) by scoring "
    "reference, Benjamini–Hochberg corrected.",
    ["Generator and evidence", "Circular reference", "Unbiased reference"],
    [["llama-3.1-8b, oracle", "+0.203", "+0.062"],
     ["gpt-oss-120b, oracle", "+0.093", "−0.021 (n.s.)"],
     ["gpt-oss-120b, real retrieval", "−0.032 (n.s.)", "−0.047"]],
)


# --------------------------------------------------------------------------
# mechanics
# --------------------------------------------------------------------------

def set_text(p, text: str) -> None:
    """Replace a paragraph's text, preserving the first run's formatting."""
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


def clone_after(p, text: str):
    """Duplicate a paragraph (keeping its style) and place the copy after it."""
    new = copy.deepcopy(p._p)
    p._p.addnext(new)
    from docx.text.paragraph import Paragraph
    np = Paragraph(new, p._parent)
    set_text(np, text)
    return np


def drop(el) -> None:
    el._element.getparent().remove(el._element)


def fill_block(doc, anchor_idx: int, paras: list[str]) -> None:
    """Replace the template paragraph at anchor_idx with one or more paragraphs."""
    p = doc.paragraphs[anchor_idx]
    set_text(p, paras[0])
    cur = p
    for extra in paras[1:]:
        cur = clone_after(cur, extra)


def build_table(doc, after_par, spec) -> None:
    caption, headers, rows = spec
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = "Times New Roman"
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
            r.font.name = "Times New Roman"
            if i:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    after_par._p.addnext(t._tbl)
    cap = clone_after(after_par, caption)
    cap.style = doc.styles["ICCSDI Caption"]


def main() -> None:
    doc = Document(str(TEMPLATE))

    # Requested: a small margin reduction, nothing else about the layout changes.
    for s in doc.sections:
        s.top_margin = Inches(0.85)
        s.bottom_margin = Inches(0.65)
        s.left_margin = Inches(0.90)
        s.right_margin = Inches(0.90)

    P = doc.paragraphs
    set_text(P[0], TITLE)
    set_text(P[1], AUTHORS)
    set_text(P[2], AFFIL)
    drop(P[3]); drop(P[4])                    # only one affiliation
    set_text(P[5], EMAILS)
    drop(P[6])                                # no third-author line
    set_text(P[8], ABSTRACT)
    set_text(P[9], KEYWORDS)

    # Re-read after deletions so indices are valid.
    def idx_of(prefix: str) -> int:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith(prefix):
                return i
        raise KeyError(prefix)

    # Body sections, filled against their template headings.
    fill_block(doc, idx_of("Explain the research context"), INTRO)
    set_text(doc.paragraphs[idx_of("A concise contribution list")],
             "The contributions are:")
    for b, txt in zip([i for i, p in enumerate(doc.paragraphs)
                       if p.style.name == "ICCSDI Bullet"][:3], INTRO_BULLETS):
        set_text(doc.paragraphs[b], "•  " + txt)
    drop(doc.paragraphs[idx_of("The extended manuscript should contain")])

    fill_block(doc, idx_of("Critically review the most relevant"), RELATED)
    fill_block(doc, idx_of("Describe the proposed method"), METHODS)
    fill_block(doc, idx_of("State the inputs, outputs"), PROBLEM)
    drop(doc.paragraphs[idx_of("where n is the number of observations")])
    fill_block(doc, idx_of("Describe all components, data flows"), ARCH)
    fill_block(doc, idx_of("Provide enough information for independent"), EXPER)
    fill_block(doc, idx_of("Describe each dataset, its provenance"), DATASETS)
    fill_block(doc, idx_of("Justify the selected baselines"), METRICS)
    fill_block(doc, idx_of("Report software libraries and versions"), IMPL)

    # Results: the template supplies a heading, one intro paragraph, a caption and
    # a follow-up. Build the five hypothesis subsections out of those, reusing the
    # template's own Heading 2 and Caption styles so nothing is invented.
    fill_block(doc, idx_of("Present findings objectively"), RES_INTRO)
    cur = doc.paragraphs[idx_of("Results are presented in hypothesis order")]

    h2 = doc.styles["ICCSDI Heading 2"]
    cap_style = doc.styles["ICCSDI Caption"]

    def add(text, style=None):
        nonlocal cur
        cur = clone_after(cur, text)
        if style is not None:
            cur.style = style
        else:
            cur.style = doc.styles["Normal"]
        return cur

    def add_table(spec):
        nonlocal cur
        caption, headers, rows = spec
        cur = clone_after(cur, caption)
        cur.style = cap_style
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(9)
            r.font.name = "Times New Roman"
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""
                r = cells[i].paragraphs[0].add_run(str(v))
                r.font.size = Pt(9)
                r.font.name = "Times New Roman"
                if i:
                    cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cur._p.addnext(t._tbl)
        cur = clone_after(cur, "")          # spacer after the table
        return t

    add("5.1  H₀₄: the retrieval-stage code-mixing penalty", h2)
    for para in RES_H4:
        add(para)
    add_table(TABLE1)

    fig = FIGDIR / "retrieval_figures" / "fig2_penalty.png"
    if fig.exists():
        run = add("").add_run()
        run.add_picture(str(fig), width=Inches(5.4))
        cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add("Fig. 1  Recall@1 lost to code-mixing by retrieval system, with 95% "
            "bootstrap intervals. Lexical retrieval is far more damaged than dense.",
            cap_style)

    add("5.2  Validity of the scoring instrument", h2)
    for para in RES_METRIC:
        add(para)

    add("5.3  H₀₁: the grounding effect", h2)
    for para in RES_H1:
        add(para)
    add_table(TABLE2)

    add("5.4  H₀₂: code-mixing robustness", h2)
    for para in RES_H2:
        add(para)

    add("5.5  H₀₃: evidence provenance", h2)
    for para in RES_H3:
        add(para)

    # The template's own example caption and follow-up are now redundant.
    for key in ("Table 1  Example comparison",
                "Replace Table 1 with verified results"):
        try:
            drop(doc.paragraphs[idx_of(key)])
        except KeyError:
            pass

    # Discussion / conclusion / limitations.
    fill_block(doc, idx_of("Interpret the results in relation"), DISCUSSION)
    fill_block(doc, idx_of("State methodological, data, measurement"), LIMITS)
    fill_block(doc, idx_of("Summarise the problem, approach"), CONCLUSION)

    # Declarations.
    for label, text in DECLARATIONS:
        i = idx_of(label)
        set_text(doc.paragraphs[i], f"{label}  {text}")

    # Supplementary / acknowledgements.
    set_text(doc.paragraphs[idx_of("State whether supplementary files")],
             "No supplementary files accompany this manuscript.")
    set_text(doc.paragraphs[idx_of("Acknowledge non-author contributions")],
             "[TO COMPLETE] Acknowledge supervisors and institutional support.")

    # Remove template scaffolding this study does not use: the worked equation,
    # the pseudocode block, the figure placeholder and the optional appendix.
    for t in list(doc.tables):
        txt = t.rows[0].cells[0].text
        if txt.startswith(("θ*", "Insert figure", "Algorithm 1", "Method")):
            drop(t)
    for key in ("3.3  Algorithm", "Provide pseudocode", "Fig. 1  Overview",
                "Appendix A", "Place material that supports"):
        try:
            drop(doc.paragraphs[idx_of(key)])
        except KeyError:
            pass

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")
    d2 = Document(str(OUT))
    print(f"  {len(d2.paragraphs)} paragraphs, {len(d2.tables)} tables")


if __name__ == "__main__":
    main()
